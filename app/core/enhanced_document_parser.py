import pypdf
from docx import Document
from typing import List, Dict, Any
from app.core.cache_manager import CacheManager
from loguru import logger
import os
import hashlib
import pathlib


class EnhancedDocumentParser:
    """
    增强的文档解析器 (支持PDF和DOCX)
    """
    
    def __init__(self, cache_manager: CacheManager = None):
        """
        初始化文档解析器
        
        Args:
            cache_manager (CacheManager, optional): 缓存管理器实例
        """
        self.cache_manager = cache_manager
        logger.info("Initialized EnhancedDocumentParser")
    
    def parse_document(self, file_path: str) -> str:
        """
        解析文档文件为文本（自动识别格式）
        
        Args:
            file_path (str): 文档文件路径
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
            Exception: 解析失败
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # 获取文件扩展名
            file_extension = pathlib.Path(file_path).suffix.lower()
            
            # 根据文件类型选择解析方法
            if file_extension == '.pdf':
                return self.parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {e}")
            raise
    
    def parse_pdf(self, file_path: str) -> str:
        """
        解析PDF文件为文本
        
        Args:
            file_path (str): PDF文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            # 生成文件路径的哈希值作为缓存键
            cache_key = f"pdf_text_{hashlib.md5(file_path.encode()).hexdigest()}"
            
            # 尝试从缓存中获取结果
            if self.cache_manager:
                cached_result = self.cache_manager.get(cache_key)
                if cached_result:
                    logger.info(f"Retrieved PDF text from cache for file: {file_path}")
                    return cached_result
            
            # 解析PDF
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += page_text + "\n"
                        logger.debug(f"Extracted text from page {page_num + 1}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            # 清理文本
            text = self._clean_text(text)
            
            # 将结果存入缓存
            if self.cache_manager:
                self.cache_manager.set(cache_key, text, expire=3600)  # 缓存1小时
                logger.info(f"Cached PDF text for file: {file_path}")
            
            logger.info(f"Parsed PDF file: {file_path}, extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse PDF file {file_path}: {e}")
            raise Exception(f"Failed to parse PDF file {file_path}: {e}")

    def parse_docx(self, file_path: str) -> str:
        """
        解析DOCX文件为文本
        
        Args:
            file_path (str): DOCX文件路径
            
        Returns:
            str: 提取的文本内容
        """
        try:
            # 生成文件路径的哈希值作为缓存键
            cache_key = f"docx_text_{hashlib.md5(file_path.encode()).hexdigest()}"
            
            # 尝试从缓存中获取结果
            if self.cache_manager:
                cached_result = self.cache_manager.get(cache_key)
                if cached_result:
                    logger.info(f"Retrieved DOCX text from cache for file: {file_path}")
                    return cached_result
            
            # 解析DOCX
            doc = Document(file_path)
            text = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # 清理文本
            text = self._clean_text(text)
            
            # 将结果存入缓存
            if self.cache_manager:
                self.cache_manager.set(cache_key, text, expire=3600)  # 缓存1小时
                logger.info(f"Cached DOCX text for file: {file_path}")
            
            logger.info(f"Parsed DOCX file: {file_path}, extracted {len(text)} characters")
            return text
            
        except Exception as e:
            logger.error(f"Failed to parse DOCX file {file_path}: {e}")
            raise Exception(f"Failed to parse DOCX file {file_path}: {e}")

    def _clean_text(self, text: str) -> str:
        """
        清理提取的文本
        
        Args:
            text (str): 原始文本
            
        Returns:
            str: 清理后的文本
        """
        # 移除多余的空白字符
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                lines.append(cleaned_line)
        
        return '\n'.join(lines)

    def parse_multiple_documents(self, file_paths: List[str]) -> Dict[str, str]:
        """
        批量解析多个文档文件
        
        Args:
            file_paths (List[str]): 文档文件路径列表
            
        Returns:
            Dict[str, str]: 文件路径到提取文本的映射
        """
        results = {}
        for file_path in file_paths:
            try:
                text = self.parse_document(file_path)
                results[file_path] = text
                logger.info(f"Successfully parsed: {file_path}")
            except Exception as e:
                logger.error(f"Failed to parse document {file_path}: {e}")
                results[file_path] = f"ERROR: {str(e)}"
        return results

    def get_supported_extensions(self) -> List[str]:
        """
        获取支持的文件扩展名列表
        
        Returns:
            List[str]: 支持的文件扩展名
        """
        return ['.pdf', '.docx', '.doc']

    def is_supported_file(self, file_path: str) -> bool:
        """
        检查文件是否为支持的格式
        
        Args:
            file_path (str): 文件路径
            
        Returns:
            bool: 是否支持该文件格式
        """
        file_extension = pathlib.Path(file_path).suffix.lower()
        return file_extension in self.get_supported_extensions()